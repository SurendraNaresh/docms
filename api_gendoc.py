from flask import Flask, request, send_file
from docx import Document
from docx.shared import Inches
from io import BytesIO
import os
app = Flask(__name__)

@app.route("/generate-doc", methods=["POST"])
def generate_doc():
    data = request.json
    doc_name = data.get("doc_name", "output")
    section_id = data.get("section_id", "Default Section")
    figure_desc = data.get("figure_desc", "")
    table_id = data.get("table_id", "")
    color_scheme = data.get("color_scheme", "0,0,0")

    # Generate Word Document
    doc = Document()
    doc.add_heading(doc_name, level=1)
    doc.add_paragraph(f"Section ID: {section_id}")
    if figure_desc and os.path.exists(figure_desc):
        doc.add_heading("Figure", level=3)
        doc.add_picture(figure_desc, width=Inches(4)) # adjust width as needed
    else:
        doc.add_paragraph(f"Image not found @{figure_desc}")
    if table_id:
        table = doc.add_table(rows=2, cols=2)
        table.style = "Light Grid Accent 1"
        table.cell(0, 0).text = "Table ID"
        table.cell(0, 1).text = table_id
   

    # Apply Color Scheme
    r, g, b = map(int, color_scheme.split(","))
    color = f"RGB({r}, {g}, {b})"  # For future use

    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(file_stream, as_attachment=True, download_name=f"{doc_name}.docx")

if __name__ == "__main__":
    app.run(port=5001)
