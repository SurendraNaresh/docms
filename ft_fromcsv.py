import csv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class CsvToWord:
    def __init__(self, csv_file, word_file):
        self.csv_file = csv_file
        self.word_file = word_file

    def read_csv(self):
        records = []
        with open(self.csv_file, mode='r') as file:
            reader = csv.reader(file)
            for row in reader:
                records.append(row)
        return records

    def write_to_word(self, records):
        doc = Document()
        for i, record in enumerate(records, start=1):dd
            # Add record number as headline
            headline = doc.add_heading(level=1)
            run = headline.add_run(f"Record {i}")
            run.bold = True

            # Add name in a tabbed paragraph section
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = paragraph.add_run(f"{record[0]} {record[1]} {record[2]}")
            run.bold = True
            run.font.size = Pt(12)

        doc.save(self.word_file)

    def convert(self):
        records = self.read_csv()
        self.write_to_word(records)

# Usage
csv_file = "names.csv"
word_file = "names.docx"
converter = CsvToWord(csv_file, word_file)
converter.convert()
